"""
Void System - User Document Parser
----------------------------------
多格式文档解析器，支持虚空系统统一文档处理流程
"""
from pathlib import Path
from typing import Dict, Any, Optional, List
import logging
from abc import ABC, abstractmethod
import pandas as pd
import io
import re

logger = logging.getLogger("void-system-doc-parser")


def _excel_sheet_display_name(sheet_name: str) -> str:
    """将 Excel 默认英文工作表名（如 Sheet1）转为中文描述，其余名称原样保留。"""
    name = str(sheet_name).strip()
    m = re.match(r"^Sheet(\d+)$", name, re.IGNORECASE)
    if m:
        return f"第 {m.group(1)} 个工作表"
    return name

class DocumentParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    def can_parse(self, file_type: str) -> bool:
        """判断是否能解析该文件类型"""
        pass

    @abstractmethod
    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析文件内容"""
        pass

    @abstractmethod
    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取文件元数据"""
        pass

class TextDocumentParser(DocumentParser):
    """文本文档解析器"""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() in ['txt', 'md', 'csv', 'json', 'py', 'js', 'html', 'css', 'xml']

    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析文本内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'utf-16', 'latin-1']
            content = None

            for encoding in encodings:
                try:
                    content = file_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                return {
                    "success": False,
                    "error": "无法解码文件内容"
                }

            # 提取文本内容
            lines = content.split('\n')
            text_content = '\n'.join(lines)

            return {
                "success": True,
                "content": text_content,
                "line_count": len(lines),
                "char_count": len(text_content),
                "language": self._detect_language(file_name, content)
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"文本解析失败: {str(e)}"
            }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取文本文件元数据"""
        try:
            content = file_data.decode('utf-8', errors='ignore')
            lines = content.split('\n')

            return {
                "file_size": len(file_data),
                "line_count": len(lines),
                "encoding": "UTF-8",
                "has_bom": file_data.startswith(b'\xef\xbb\xbf'),
                "language": self._detect_language(file_name, content[:1000])  # 只检测前1000字符
            }
        except Exception as e:
            return {
                "file_size": len(file_data),
                "error": f"元数据提取失败: {str(e)}"
            }

    def _detect_language(self, file_name: str, content: str) -> str:
        """简单语言检测"""
        ext = file_name.split('.')[-1].lower()

        # 基于扩展名
        lang_map = {
            'py': 'python',
            'js': 'javascript',
            'ts': 'typescript',
            'java': 'java',
            'cpp': 'cpp',
            'c': 'c',
            'cs': 'csharp',
            'php': 'php',
            'rb': 'ruby',
            'go': 'go',
            'rs': 'rust',
            'md': 'markdown',
            'html': 'html',
            'css': 'css',
            'xml': 'xml',
            'json': 'json',
            'yaml': 'yaml',
            'yml': 'yaml',
            'sql': 'sql',
            'sh': 'shell',
            'txt': 'text'
        }

        return lang_map.get(ext, 'unknown')

class PDFDocumentParser(DocumentParser):
    """PDF文档解析器"""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == 'pdf'

    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析PDF内容"""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(file_data)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for i in range(len(reader.pages)):
                page = reader.pages[i]
                text_parts.append(page.extract_text())
                
            content = "\n\n".join(text_parts)
            
            return {
                "success": True,
                "content": content,
                "page_count": len(reader.pages),
                "char_count": len(content)
            }
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            return {
                "success": False,
                "error": f"PDF解析失败: {str(e)}"
            }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取PDF元数据"""
        return {
            "file_size": len(file_data),
            "pages": "unknown",  # 需要PDF库支持
            "has_text": True  # 假设有文本
        }

class WordDocumentParser(DocumentParser):
    """Word文档解析器"""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() in ['doc', 'docx']

    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析Word内容"""
        try:
            import docx
            import io
            
            doc_file = io.BytesIO(file_data)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 也可以处理表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
            content = "\n\n".join(text_parts)
            
            return {
                "success": True,
                "content": content,
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(content)
            }
        except Exception as e:
            logger.error(f"Word解析失败: {e}")
            return {
                "success": False,
                "error": f"Word解析失败: {str(e)}"
            }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取Word元数据"""
        return {
            "file_size": len(file_data),
            "pages": "unknown",
            "has_text": True
        }

class ExcelDocumentParser(DocumentParser):
    """Excel文档解析器"""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() in ['xls', 'xlsx', 'csv']

    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析Excel内容"""
        try:
            if file_name.lower().endswith('.csv'):
                # CSV处理 - 使用pandas优化
                try:
                    # 使用pandas读取CSV
                    df = pd.read_csv(io.BytesIO(file_data), encoding='utf-8', on_bad_lines='skip')
                    content = self._dataframe_to_text(df, "CSV数据")
                    return {
                        "success": True,
                        "content": content,
                        "row_count": len(df),
                        "column_count": len(df.columns),
                        "format": "csv"
                    }
                except UnicodeDecodeError:
                    # 回退到原始方法
                    content = file_data.decode('gbk', errors='ignore')
                    lines = content.split('\n')
                    return {
                        "success": True,
                        "content": content,
                        "row_count": len([l for l in lines if l.strip()]),
                        "format": "csv"
                    }
            else:
                # Excel文件处理 (.xls, .xlsx)
                try:
                    # 使用pandas读取Excel
                    excel_data = pd.ExcelFile(io.BytesIO(file_data))

                    # 处理所有工作表
                    all_content = []
                    total_rows = 0
                    total_columns = 0

                    for sheet_name in excel_data.sheet_names:
                        df = pd.read_excel(excel_data, sheet_name=sheet_name)
                        sheet_content = self._dataframe_to_text(
                            df, _excel_sheet_display_name(sheet_name)
                        )
                        all_content.append(sheet_content)
                        total_rows += len(df)
                        total_columns = max(total_columns, len(df.columns))

                    content = "\n\n".join(all_content)

                    return {
                        "success": True,
                        "content": content,
                        "sheet_count": len(excel_data.sheet_names),
                        "total_rows": total_rows,
                        "total_columns": total_columns,
                        "sheet_names": excel_data.sheet_names,
                        "format": file_name.split('.')[-1].lower()
                    }

                except Exception as excel_error:
                    return {
                        "success": False,
                        "error": f"Excel文件解析失败: {str(excel_error)}。请确保文件格式正确且未损坏。"
                    }

        except Exception as e:
            return {
                "success": False,
                "error": f"文件解析失败: {str(e)}"
            }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取Excel元数据"""
        try:
            # 尝试读取Excel文件获取更多元数据
            excel_data = pd.ExcelFile(io.BytesIO(file_data))

            metadata = {
                "file_size": len(file_data),
                "format": file_name.split('.')[-1].lower(),
                "has_data": len(file_data) > 0,
                "sheet_count": len(excel_data.sheet_names),
                "sheet_names": excel_data.sheet_names
            }

            # 获取第一个工作表的统计信息
            if excel_data.sheet_names:
                df = pd.read_excel(excel_data, sheet_name=0, nrows=0)  # 只读取表头
                metadata.update({
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist()
                })

            return metadata

        except Exception as e:
            # 如果读取失败，返回基本信息
            return {
                "file_size": len(file_data),
                "format": file_name.split('.')[-1].lower(),
                "has_data": len(file_data) > 0,
                "error": f"元数据提取失败: {str(e)}"
            }

    def _dataframe_to_text(self, df: pd.DataFrame, title: str) -> str:
        """
        将pandas DataFrame转换为适合向量化处理的文本格式
        增强语义描述，使每一行数据在向量空间中更具辨识度
        """
        if df.empty:
            return f"表格: {title}\n(数据内容为空)"

        # 基础列信息
        cols = df.columns.tolist()
        
        # 构建增强语料
        lines = [f"### 数据表: {title}"]
        lines.append(f"该表包含 {len(df)} 条记录，关键列包括: {', '.join(cols)}")
        lines.append("--- 记录详情 ---")

        for idx, row in df.iterrows():
            # 为每一行生成一个具有独立语境的描述
            row_desc = [f"记录第 {idx+1} 号:"]
            for col in cols:
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    # 采用 "[字段名] 是 [值]" 的自然语言结构，提升嵌入模型的理解力
                    row_desc.append(f"{col} 是 \"{val}\"")
            
            if len(row_desc) > 1:
                lines.append(" ".join(row_desc))
            
            # 性能限制
            if idx >= 1500: # 稍微放宽限制
                lines.append(f"... (后续 {len(df)-1500} 行数据已省略)")
                break

        return "\n".join(lines)

class ImageDocumentParser(DocumentParser):
    """图片文档解析器（OCR）"""

    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']

    def parse_content(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """解析图片内容（OCR）"""
        try:
            # 这里需要安装pytesseract和PIL
            return {
                "success": False,
                "error": "OCR解析器暂未实现，请安装pytesseract和Pillow库"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"图片解析失败: {str(e)}"
            }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """提取图片元数据"""
        return {
            "file_size": len(file_data),
            "has_ocr": False,  # 暂不支持OCR
            "image_type": file_name.split('.')[-1].lower()
        }

class UniversalDocumentParser:
    """通用文档解析器"""

    def __init__(self):
        self.parsers = [
            TextDocumentParser(),
            PDFDocumentParser(),
            WordDocumentParser(),
            ExcelDocumentParser(),
            ImageDocumentParser()
        ]

    def parse_file(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """
        解析文件内容
        Args:
            file_data: 文件数据
            file_name: 文件名
        Returns:
            解析结果
        """
        file_type = self._get_file_type(file_name)

        # 查找合适的解析器
        for parser in self.parsers:
            if parser.can_parse(file_type):
                result = parser.parse_content(file_data, file_name)
                result["file_type"] = file_type
                return result

        # 没有找到合适的解析器
        return {
            "success": False,
            "error": f"不支持的文件类型: {file_type}",
            "file_type": file_type
        }

    def extract_metadata(self, file_data: bytes, file_name: str) -> Dict[str, Any]:
        """
        提取文件元数据
        Args:
            file_data: 文件数据
            file_name: 文件名
        Returns:
            元数据字典
        """
        file_type = self._get_file_type(file_name)
        base_metadata = {
            "file_name": file_name,
            "file_type": file_type,
            "file_size": len(file_data),
            "file_extension": file_name.split('.')[-1].lower() if '.' in file_name else ''
        }

        # 查找合适的解析器
        for parser in self.parsers:
            if parser.can_parse(file_type):
                metadata = parser.extract_metadata(file_data, file_name)
                base_metadata.update(metadata)
                return base_metadata

        return base_metadata

    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型"""
        supported = []
        for parser in self.parsers:
            # 这里需要实现一个方法来获取每个解析器支持的类型
            # 暂时返回常用类型
            pass
        return ['txt', 'md', 'pdf', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'jpg', 'png']

    def _get_file_type(self, file_name: str) -> str:
        """从文件名获取文件类型"""
        if '.' not in file_name:
            return 'unknown'

        ext = file_name.split('.')[-1].lower()

        # 扩展名到类型映射
        type_map = {
            'txt': 'txt', 'md': 'md', 'markdown': 'md',
            'pdf': 'pdf',
            'doc': 'doc', 'docx': 'docx',
            'xls': 'xls', 'xlsx': 'xlsx', 'csv': 'csv',
            'jpg': 'jpg', 'jpeg': 'jpeg', 'png': 'png',
            'gif': 'gif', 'bmp': 'bmp', 'tiff': 'tiff',
            'py': 'py', 'js': 'js', 'html': 'html', 'css': 'css',
            'json': 'json', 'xml': 'xml', 'yaml': 'yaml', 'yml': 'yaml'
        }

        return type_map.get(ext, 'unknown')

# 全局解析器实例
document_parser = UniversalDocumentParser()

